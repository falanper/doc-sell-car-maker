print('='*46)
print('     GERADOR DE TERMO DE VENDA DE VEÍCULO')
print('='*46)
print('Por favor, responda as seguintes perguntas:')
comprador = str(input('\nDigite o nome completo do comprador: ')).strip()
sexo_comprador = sexo_vendedor = meio_pagamento = ' '
filler_comprador = filler_vendedor = ' '
month = 0
while sexo_comprador not in 'MmFf':
    sexo_comprador = str(input('Digite o sexo do comprador [M/F]: ')).strip()
if sexo_comprador in 'Mm':
    artigo_comprador = 'o'
    tratamento_comprador = 'Sr.'
    preposicao_comprador = 'ao'
    filler_comprador = ' '
if sexo_comprador in 'Ff':
    artigo_comprador = 'a'
    tratamento_comprador = 'Sra.'
    preposicao_comprador = 'à'
    filler_comprador = 'a'
cpf_comprador = str(input('Digite o CPF do comprador [formato XXX.XXX.XXX-XX]: ')).strip()
vendedor = str(input('Digite o nome completo do vendedor: ')).strip()
while sexo_vendedor not in 'MmFf':
    sexo_vendedor = str(input('Digite o sexo do vendedor [M/F]: ')).strip()
if sexo_vendedor in 'Mm':
    artigo_vendedor = 'o'
    tratamento_vendedor = 'Sr.'
    preposicao_vendedor = 'ao'
    filler_vendedor = ' '
if sexo_vendedor in 'Ff':
    artigo_vendedor = 'a'
    tratamento_vendedor = 'Sra.'
    preposicao_vendedor = 'à'
    filler_vendedor = 'a'
cpf_vendedor = str(input('Digite o CPF do vendedor [formato XXX.XXX.XXX-XX]: ')).strip()
marca = str(input('Digite a marca do veículo: ')).strip()
modelo = str(input('Digite o modelo do veículo: ')).strip()
ano = str(input('Digite o ano do veículo: ')).strip()
placa = str(input('Digite a placa do veículo: '))
valor = str(input('Digite o valor de venda [sem "R$"]: ')).strip()
while meio_pagamento not in 'TtDd':
    meio_pagamento = str(input('Digite o meio de pagamento ["T" para TED, "D" para dinheiro: '))
if meio_pagamento in 'Tt':
    conta = str(input('Digite o número da conta corrente do vendedor: ')).strip()
    agencia = str(input('Digite o número da agência do vendedor: ')).strip()
    banco = str(input('Digite o nome do banco do vendedor [sem a palavra "Banco"]: ')).strip()
    meio_pagamento = f'transferência bancária realizada nesta data para a Conta Corrente n. {conta}, ' \
        f'mantida na Agência n. {agencia}, do Banco {banco}, de titularidade de {vendedor}'
if meio_pagamento in 'Dd':
    meio_pagamento = 'pagamento em dinheiro realizado nesta data'
city = str(input('Digite o nome da cidade onde será realizada a venda: ')).strip()
day = str(input('Digite o dia da venda: ')).strip()
month_range = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]
while month not in month_range:
    try:
        month = int(input('Digite o mês da venda [digite entre "1" e "12"]: '))
    except ValueError:
        print('Infelizmente o número digitado não está entre 1 e 12. Por favor tente novamente.')
year = str(input('Digite o ano da venda: ')).strip()
months = ('zero', 'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho', 'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro')
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
heading = 'TERMO DE RESPONSABILIDADE POR VENDA DE VEÍCULO'
doc = docx.Document()
p = doc.add_paragraph()
p.add_run(heading).bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def pulalinha():
    doc.add_paragraph()


pulalinha()
doc.save('Termo de Responsabilidade.docx')
paragraph1 = doc.add_paragraph(f'{artigo_vendedor.upper()} vendedor{filler_vendedor} {tratamento_vendedor} {vendedor} delara haver recebido d{artigo_comprador} {tratamento_comprador} {comprador},'
                               f' inscrit{artigo_comprador} no CPF/MF sob o n.º {cpf_comprador}, o valor total de R$ {valor},'
                               f' relativos à venda do veículo {marca} {modelo}, Ano {ano}, Placa {placa},'
                               f' pago em moeda corrente nacional, por meio de {meio_pagamento},'
                               f' pagamento do qual {artigo_vendedor} vendedor{filler_vendedor} dá total quitação {preposicao_comprador} {tratamento_comprador} {comprador}.')
paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pulalinha()
paragraph2 = doc.add_paragraph(f'{artigo_comprador.upper()} {tratamento_comprador} {comprador}, de posse do veículo acima referido'
                               f' a partir das ___:___ horas deste dia {day} de {months[month]} de {year}, se responsabiliza integralmente'
                               f' por quaisquer atos praticados com o veículo ou fatos a ele relacionados a partir do dia e hora'
                               f' acima referidos, isentando integralmente {artigo_vendedor} vendedor{filler_vendedor} de qualquer responsabilidade pelo veículo,'
                               f' ressalvada a responsabilidade d{artigo_vendedor} vendedor{filler_vendedor} por atos ou fatos anteriores aos referidos dia e hora,'
                               f' inclusive eventuais multas e débitos passados do veículo.')
paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pulalinha()
paragraph3 = doc.add_paragraph(f'{city}, {day} de {months[month]} de {year}.')
paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('Termo de Responsabilidade.docx')
pulalinha()
v = vendedor
table = doc.add_table(rows=4, cols=2)
row1 = table.rows[0].cells
row1[0].text = ('_'*35)
row1[1].text = ('_'*35)
row2 = table.rows[1].cells
row2[0].text = (f'{vendedor}')
row2[1].text = (f'{comprador}')
row3 = table.rows[2].cells
row3[0].text = (f'{cpf_vendedor}')
row3[1].text = (f'{cpf_comprador}')
row4 = table.rows[3].cells
row4[0].text = (f'Vendedor{filler_vendedor}')
row4[1].text = (f'Comprador{filler_comprador}')
doc.save('Termo de Responsabilidade.docx')
print('='*46)
print('Seu documento está pronto! \nVeja como ele ficou acessando na pasta com o nome "Termo de Responsabilidade".')
print('\nSeu programa será encerrado automaticamente em alguns segundos.')
import time
time.sleep(8)
